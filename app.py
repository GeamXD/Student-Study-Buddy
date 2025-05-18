import streamlit as st
from langchain.prompts import PromptTemplate
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.output_parsers import PydanticOutputParser
from pydantic import BaseModel, Field
from PyPDF2 import PdfReader
import docx
from langchain.schema import Document
from langchain_core.callbacks.base import BaseCallbackHandler
import io
from bson.objectid import ObjectId

from agent import get_agent_executor, get_search_tool
from ext_tools.qa_tool import qa_generation
from ext_tools.instant_rag import create_rag_tool

from utils.database import (
    get_chat_sessions,
    create_chat_session,
    prepare_chat_history,
    add_message_to_session,
    update_session_name
)

class LambdaStreamlitLoader:
    def __init__(self, uploaded_file) -> None:
        self.uploaded_file = uploaded_file
        self.file_name = uploaded_file.name
        if "." in self.file_name:
            *_, self.ext = self.file_name.rsplit(".", 1)
            self.ext = self.ext.lower()
        else:
            self.ext = ""

    def lazy_load(self):
        """Yields Document objects line by line from the uploaded file."""
        try:
            file_content = self.uploaded_file.getvalue()
            if not file_content:
                 st.warning(f"File '{self.file_name}' appears to be empty.")
                 return

            if self.ext in ["docx", "doc"]:
                doc = docx.Document(io.BytesIO(file_content))
                for paragraph in doc.paragraphs:
                    lines = paragraph.text.split("\n")
                    for line in lines:
                        line = line.strip()
                        if line:
                            yield Document(page_content=line, metadata={"source": self.file_name})

            elif self.ext == "pdf":
                reader = PdfReader(io.BytesIO(file_content))
                if not reader.pages:
                    st.warning(f"Could not read any pages from PDF '{self.file_name}'. It might be empty or corrupted.")
                    return
                for i, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text:
                        lines = text.split("\n")
                        for line in lines:
                            line = line.strip()
                            if line:
                                yield Document(page_content=line, metadata={"source": self.file_name, "page": i + 1})
            else:
                st.warning(f"Unsupported file format: '{self.ext if self.ext else 'Unknown'}'. Only PDF or DOCX are processed for context.")
                return

        except Exception as e:
            st.error(f"Error processing file '{self.file_name}': {e}")
            return

def get_base_title(unique_title: str) -> str:
    """Extracts the base title from a unique title (title_sessionid)."""
    if not unique_title or not isinstance(unique_title, str):
        return "Chat"
    parts = unique_title.rsplit('_', 1)
    if len(parts) == 2 and len(parts[1]) == 24:
        try:
            ObjectId(parts[1])
            return parts[0]
        except Exception:
            return unique_title
    else:
        return unique_title

class TitleParser(BaseModel):
    title: str = Field(description="Title of the chat session")

title_parser = PydanticOutputParser(pydantic_object=TitleParser)

def generate_title_llm(first_message: str) -> str:
    """Generates the base title (without unique ID)."""
    prompt = PromptTemplate(
        template="Based on the given message, suggest a suitable title for the chat (max 5 words).\nMessage: {message}\n{format_instructions}",
        input_variables=["message"],
        partial_variables={"format_instructions": title_parser.get_format_instructions()}
    )
    llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.2)
    chain = prompt | llm | title_parser
    try:
        result = chain.invoke({"message": first_message})
        title = result.title
        title = title.strip().strip('"')
        return title if title else "Chat Session"
    except Exception as e:
        st.error(f"Title generation failed: {e}")
        return "Chat Session"

class StreamlitCallbackHandler(BaseCallbackHandler):
    def __init__(self, status):
        self.status = status
        self.action = None

    def on_agent_action(self, action, **kwargs):
        self.action = action
        tool_input = action.tool_input
        query = "details"
        tool_name = action.tool

        if isinstance(tool_input, dict):
            query = tool_input.get("query", tool_input.get("input", str(tool_input)))
        elif isinstance(tool_input, str):
             query = tool_input

        if tool_name == "document_search":
             self.status.update(label=f"Searching document for: `{query}`")
        elif tool_name == "tavily_search_results_json":
             self.status.update(label=f"Searching the web for: `{query}`")
        elif tool_name == "qa_generation":
             num_pairs = tool_input.get('number', 'some') if isinstance(tool_input, dict) else 'some'
             self.status.update(label=f"Generating {num_pairs} Q&A pairs...")
        else:
             self.status.update(label=f"Using tool `{tool_name}`...")


    def on_tool_start(self, serialized, input_str, **kwargs):
        tool_name = serialized.get("name", "tool")
        query = "details"
        if self.action:
            tool_input = self.action.tool_input
            if isinstance(tool_input, dict):
                query = tool_input.get("query", tool_input.get("input", str(tool_input)))
            elif isinstance(tool_input, str):
                query = tool_input

        if tool_name == "document_search":
            self.status.update(label=f"Processing document search for: `{query}`")
        elif tool_name == "tavily_search_results_json":
            self.status.update(label=f"Processing web search for: `{query}`")
        elif tool_name == "qa_generation":
            self.status.update(label=f"Processing Q&A generation...")
        else:
            self.status.update(label=f"Processing with `{tool_name}`...")


    def on_tool_end(self, output, **kwargs):
        tool_name = self.action.tool if self.action else "tool"
        if tool_name == "document_search":
            self.status.update(label="Aggregating document search results")
        elif tool_name == "tavily_search_results_json":
            self.status.update(label="Aggregating web search results")
        elif tool_name == "qa_generation":
            self.status.update(label="Finalizing Q&A generation")
        else:
            self.status.update(label=f"Finished using `{tool_name}`")

def initialize_session_state():
    default_session_state = {
        "logged_in": False,
        "email": "",
        "username": "",
        "current_session_id": None,
        "current_session_title": "",
        "needs_title": False,
        "session_selected": False,
        "processed_file_id": None,
        "file_docs": None,
        "tools": [get_search_tool()],
        "downloadable_csv": None
    }
    for key, default_value in default_session_state.items():
        if key not in st.session_state:
            st.session_state[key] = default_value

initialize_session_state()

if not st.experimental_user.is_logged_in:
    col1, col_main, col3 = st.columns([1, 5, 1])

    with col_main:
        st.header(f"Agent Lambda: AI Assistant", anchor=False)

        with st.container(border=True):
            st.markdown("<br>", unsafe_allow_html=True)
            google_logo_url = "https://upload.wikimedia.org/wikipedia/commons/c/c1/Google_%22G%22_logo.svg"
            st.markdown(
                f"""
                <div style="display: flex; justify-content: center; margin-top: 20px; margin-bottom: 20px;">
                    <img src="{google_logo_url}" width="25" height="25" style="margin-right: 10px; vertical-align: middle;">
                </div>
                """,
                unsafe_allow_html=True,
            )

            btn_col1, btn_col2, btn_col3 = st.columns([1, 2, 1])
            with btn_col2:
                login_pressed = st.button("Continue with Google", use_container_width=True)
                if login_pressed:
                    try:
                        st.login()
                    except Exception as e:
                        st.error(f"Login setup error: {e}. Ensure Google OAuth is configured.")

            st.markdown("<br>", unsafe_allow_html=True)

    st.stop()
else:
    if not st.session_state.logged_in or st.session_state.email != st.experimental_user.email:
        st.session_state.logged_in = True
        st.session_state.email = st.experimental_user.email
        st.session_state.username = st.experimental_user.name
        st.session_state.current_session_id = None
        st.session_state.current_session_title = ""
        st.session_state.needs_title = False
        st.session_state.session_selected = False
        st.session_state.processed_file_id = None
        st.session_state.file_docs = None
        st.session_state.tools = [get_search_tool()]
        st.session_state.downloadable_csv = None


with st.sidebar:
    st.write(f"Welcome, {st.session_state.username}")

    if st.button("➕ Create New Chat"):
        try:
            new_session_id, unique_name = create_chat_session(user_id=st.session_state.email, session_name="New Chat")
            st.session_state.current_session_id = new_session_id
            st.session_state.current_session_title = unique_name
            st.session_state.needs_title = True
            st.session_state.session_selected = True
            st.session_state.processed_file_id = None
            st.session_state.file_docs = None
            st.session_state.tools = [get_search_tool()]
            st.session_state.downloadable_csv = None
            st.success("New chat created. Send your first message!")
            st.rerun()
        except Exception as e:
            st.error(f"Failed to create new chat: {e}")


    st.divider()
    if st.session_state.session_selected:
        uploader_key = f"file_uploader_{st.session_state.current_session_id}" if st.session_state.current_session_id else "file_uploader_default"
        uploaded_file = st.file_uploader(
            "Upload PDF or DOCX (Optional)",
            type=["pdf", "docx"],
            key=uploader_key
        )

        if uploaded_file is not None and uploaded_file.file_id != st.session_state.get("processed_file_id"):
            try:
                with st.spinner(f"Processing file: {uploaded_file.name}..."):
                    loader = LambdaStreamlitLoader(uploaded_file)
                    docs = list(loader.lazy_load())

                    if docs:
                        st.session_state.file_docs = docs
                        st.session_state.processed_file_id = uploaded_file.file_id
                        st.success(f"File '{uploaded_file.name}' uploaded successfully! Preparing document tools...")

                        current_tools = [get_search_tool()]

                        rag_tool = create_rag_tool(st.session_state.file_docs)
                        if rag_tool:
                            current_tools.append(rag_tool)

                        qa_tool = qa_generation
                        current_tools.append(qa_tool)

                        st.session_state.tools = current_tools
                        st.session_state.downloadable_csv = None

                        st.rerun()

                    else:
                        st.warning(f"Could not extract content from '{uploaded_file.name}'. Tools requiring document context will not be added.")
                        st.session_state.processed_file_id = uploaded_file.file_id
                        st.session_state.file_docs = None
                        st.session_state.tools = [get_search_tool()]
                        st.session_state.downloadable_csv = None
                        st.rerun()

            except Exception as e:
                st.error(f"Failed to process file: {e}")
                st.session_state.processed_file_id = uploaded_file.file_id
                st.session_state.file_docs = None
                st.session_state.tools = [get_search_tool()]
                st.session_state.downloadable_csv = None
                st.rerun()

        if st.session_state.get("file_docs"):
            st.info(f"File loaded. Document tools active.")
        elif st.session_state.get("processed_file_id") and not st.session_state.get("file_docs"):
             st.warning("A file was processed, but no content was loaded or an error occurred.")
    else:
        st.info("Select or create a chat to enable file upload.")
    

    st.divider()
    try:
        sessions_list = list(get_chat_sessions(user_id=st.session_state.email))
    except Exception as e:
        st.error(f"Failed to load chat sessions: {e}")
        sessions_list = []

    if sessions_list:
        session_options = {unique_name: sid for sid, unique_name in sessions_list}
        session_unique_names = [unique_name for sid, unique_name in sessions_list]

        current_index = 0
        current_title_str = str(st.session_state.get("current_session_title", ""))
        if st.session_state.current_session_id and current_title_str:
            try:
                current_index = session_unique_names.index(current_title_str)
            except ValueError:
                if session_unique_names:
                    first_unique_name = session_unique_names[0]
                    st.session_state.current_session_id = session_options[first_unique_name]
                    st.session_state.current_session_title = first_unique_name
                    st.session_state.needs_title = False
                    st.session_state.session_selected = True
                    st.session_state.processed_file_id = None
                    st.session_state.file_docs = None
                    st.session_state.tools = [get_search_tool()]
                    st.session_state.downloadable_csv = None
                    current_index = 0
                    st.rerun()
                else:
                    st.session_state.current_session_id = None
                    st.session_state.current_session_title = ""
                    st.session_state.needs_title = False
                    st.session_state.session_selected = False
                    st.session_state.processed_file_id = None
                    st.session_state.file_docs = None
                    st.session_state.tools = [get_search_tool()]
                    st.session_state.downloadable_csv = None

        if session_unique_names:
            selected_session_name = st.selectbox(
                "Chat History",
                options=session_unique_names,
                index=current_index,
                key="session_selector",
                format_func=get_base_title
            )

            selected_session_id = session_options.get(selected_session_name)
            if selected_session_id and selected_session_id != st.session_state.current_session_id:
                st.session_state.current_session_id = selected_session_id
                st.session_state.current_session_title = selected_session_name
                st.session_state.needs_title = False
                st.session_state.session_selected = True
                st.session_state.processed_file_id = None
                st.session_state.file_docs = None
                st.session_state.tools = [get_search_tool()]
                st.session_state.downloadable_csv = None
                st.rerun()

    else:
        st.warning("No previous chats found. Create one!")
    


if st.session_state.logged_in and st.session_state.session_selected and st.session_state.current_session_id:
    current_title_str = str(st.session_state.get("current_session_title", ""))
    display_title = get_base_title(current_title_str)
    st.subheader(f"Chat: {display_title}", anchor=False)

    try:
        session_id_obj = st.session_state.current_session_id
        if not isinstance(session_id_obj, ObjectId):
             session_id_obj = ObjectId(session_id_obj)

        messages = prepare_chat_history(
            session_id=session_id_obj,
            chat_history_limit=50
        )
    except Exception as e:
        st.error(f"Failed to load chat history: {e}")
        messages = []

    for message in messages:
        role = getattr(message, 'kind', getattr(message, 'type', 'ai' if hasattr(message, 'ai') else 'human'))
        display_role = "user" if role in ["human", "user"] else "assistant"
        with st.chat_message(display_role):
             st.markdown(str(message.content))

    if st.session_state.get('downloadable_csv'):
        csv_data = st.session_state['downloadable_csv']
        st.download_button(
            label="Download QA Pairs as CSV",
            data=csv_data["content"],
            file_name=csv_data["filename"],
            mime=csv_data["mime"],
            key=f"download_{csv_data['filename']}"
        )
        st.session_state.downloadable_csv = None


    prompt_text = st.chat_input("Ask your questions...")
    if prompt_text:
        st.chat_message("user").markdown(prompt_text)

        session_id_to_use = st.session_state.current_session_id
        if not isinstance(session_id_to_use, ObjectId):
            try:
                session_id_to_use = ObjectId(session_id_to_use)
            except Exception as e:
                st.error(f"Internal Error: Invalid session ID format during title generation. {e}")
                st.stop()

        title_generated_in_this_run = False
        if st.session_state.needs_title:
            try:
                with st.spinner("Generating title"):
                    base_title = generate_title_llm(prompt_text)
                new_unique_title = update_session_name(session_id_to_use, base_title)
                if new_unique_title:
                    st.session_state.current_session_title = new_unique_title
                    st.session_state.needs_title = False
                    title_generated_in_this_run = True
                else:
                     st.warning("Failed to update title in database, keeping old title.")
                     st.session_state.needs_title = False

            except Exception as title_e:
                st.error(f"Failed to generate or update title: {title_e}")
                st.session_state.needs_title = False


        try:
            add_message_to_session(
                session_id=session_id_to_use,
                content=prompt_text,
                kind="user"
            )
        except Exception as e:
             st.error(f"Failed to save your message: {e}")

        try:
            messages = prepare_chat_history(
                session_id=session_id_to_use,
                chat_history_limit=50
            )
        except Exception as e:
            st.error(f"Failed to reload chat history: {e}")


        agent_input = {"input": prompt_text, "chat_history": messages}

        current_tools = st.session_state.get("tools", [get_search_tool()])
        try:
            agent_executor = get_agent_executor(tools=current_tools)
        except Exception as agent_init_e:
             st.error(f"Failed to initialize the AI agent: {agent_init_e}")
             st.stop()


        with st.status("Processing your request...", expanded=False) as status:
            try:
                callback_handler = StreamlitCallbackHandler(status)
                response = agent_executor.invoke(
                    agent_input,
                    config={"callbacks": [callback_handler]}
                )

                output = response.get("output", "Sorry, I couldn't process that.")
                status.update(label="Done!", state="complete", expanded=True)
                with st.chat_message("assistant"):
                    st.markdown(str(output))
                    
                try:
                    add_message_to_session(
                        session_id=session_id_to_use,
                        content=str(output),
                        kind="ai"
                    )
                except Exception as e:
                    st.error(f"Failed to save AI response: {e}")

                st.rerun()
            except Exception as e:
                error_message = f"An error occurred while processing your request: {e}. Please try again."
                if "rate limit" in str(e).lower() or "429" in str(e):
                     error_message = "Apologies, the system is experiencing high load (rate limit exceeded). Please try again in a few moments."
                elif "API key not valid" in str(e):
                     error_message = "Configuration error: An API key is invalid. Please contact support."


                status.update(label=f"Error: Processing failed.", state="error", expanded=True)
                st.chat_message("assistant").error(error_message)